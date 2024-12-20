from summarizer import Summarizer
import docx
import os
from dotenv import load_dotenv

together_api_key = os.getenv("TOGETHER_API_KEY")

def read_text_from_docx(file_path):
    """
    Read text from a Word document.
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        return '\n'.join(paragraphs)
    except Exception as e:
        print(f"Error reading document: {e}")
        return None


def summarize_chunk(summarizer, chunk):
    """
    Generate a summary for a chunk of text.
    """
    return summarizer.generate_summary(chunk)


def generate_biography(summarizer, summary):
    """
    Generate a summary for a chunk of text.
    """
    return summarizer.generate_biography(summary)

def improve_coherence(summarizer, summary):
    """
    Improve the coherence of the biography.
    """
    return summarizer.improve_coherence(summary)


def divide_into_chunks(text, word_per_chunk):
    """
    Divide input text into chunks based on word count.
    """
    chunks = []
    words = text.split()
    num_chunks = len(words) // word_per_chunk
    for i in range(num_chunks):
        start = i * word_per_chunk
        end = (i + 1) * word_per_chunk
        chunks.append(" ".join(words[start:end]))
    # Adding the remaining words as the last chunk
    if len(words) % word_per_chunk != 0:
        chunks.append(" ".join(words[num_chunks * word_per_chunk:]))
    return chunks


def remove_incomplete_sentence(biography):
    """
    Remove incomplete sentence at the end of the biography.
    """
    last_full_stop_index = biography.rfind('.')
    
    if last_full_stop_index != -1:
        return biography[:last_full_stop_index + 1]
    else:
        return biography



# def biography(file_path, word_per_chunk=1000, output_file="summarized_document.docx"):
def biography(input_text, word_per_chunk=1000, output_file="summarized_document.docx"):

    """
    Summarize a document and save the summary to a new document.
    """
    # input_text = read_text_from_docx(file_path)
    # if input_text is None:
    #     return

    summarizer = Summarizer()
    biography = input_text
    while True:
        chunks = divide_into_chunks(biography, word_per_chunk)
        summary = ""
        for chunk in chunks:
            summary += summarize_chunk(summarizer, chunk) + "\n"
        total_tokens = len(summary.split())
        if total_tokens <= 600: #mistralai/Mixtral-8x7B-Instruct-v0.1, togethercomputer/StripedHyena-Nous-7B 
        # if total_tokens > 490: #snorkelai/Snorkel-Mistral-PairRM-DPO, mistralai/Mistral-7B-Instruct-v0.2
        # if total_tokens > 400: #upstage/SOLAR-10.7B-Instruct-v1.0
            break
        biography = summary

    summary = generate_biography(summarizer, summary)
    # summary = improve_coherence(summarizer, summary)
    summary = remove_incomplete_sentence(summary)
    # new_doc = docx.Document()
    # new_doc.add_paragraph(summary)
    # new_doc.save(output_file)
    return summary


if __name__ == "__main__":
    document_path = "Transkript_sample.docx"
    biography(document_path)
