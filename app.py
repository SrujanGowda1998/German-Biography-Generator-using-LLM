# import streamlit as st
# from document_reader_v5 import read_text_from_docx, biography
# from docx import Document
# import io

# # This must be the first Streamlit command used.
# st.set_page_config(page_title="German Biography Generator", layout="wide")

# # Custom header with HTML and CSS
# header_html = """
#     <div style="background-color:#618264; padding:10px; border-radius:10px">
#     <h1 style="color:white; text-align:center;">German Biography Generator</h1>
#     <p style="color:white; text-align:center;">We generate short biographies in German language using interview text of any length. Please upload the interview text in Word format to generate biography of the interviewee.</p>
#     </div>
#     """
# st.markdown(header_html, unsafe_allow_html=True)

# # Inject CSS for button color and footer positioning
# st.markdown("""
# <style>
# .stButton>button {
#     border: 1px solid #4CAF50;
#     border-radius: 4px; /* Match Streamlit's default button radius */
#     color: black; /* Match Streamlit's default button text color */
#     background-color: #D0E7D2; /* Custom button color */
#     font-size: 16px; /* Adjust to match Streamlit's default button size, if needed */
# }

# footer {
#     position: bottom;
#     bottom: 0;
#     background-color: #333;
#     color: #FFFFFF;
#     text-align: center;
#     padding: 10px 0;
#     margin-top: 50px;
#     border-radius: 10px;
# }

# footer a {
#     color: #FFFFFF;
#     left: 0;
#     text-decoration: none;
# }

# footer a:hover {
#     text-decoration: underline;
# }
# </style>
# """, unsafe_allow_html=True)

# # Footer HTML
# footer_html = """
# <footer>
#     <p style="font-size: 14px;">Contact developer: <a href="mailto:srujanprakashgowda3@gmail.com">srujanprakashgowda3@gmail.com</a></p>
#     <p style="font-size: 14px;">
#         <a href="https://www.linkedin.com/in/srujanpg/" target="_blank">LinkedIn</a> |
#         <a href="https://github.com/SrujanGowda1998" target="_blank">GitHub</a>
#     </p>
# </footer>
# """

# def create_word_document(biography_text):
#     doc = Document()
#     doc.add_paragraph(biography_text)
#     doc_io = io.BytesIO()
#     doc.save(doc_io)
#     doc_io.seek(0)
#     return doc_io

# if 'biography_generated' not in st.session_state:
#     st.session_state.biography_generated = False

# if not st.session_state.biography_generated:
#     uploaded_file = st.file_uploader("Choose a Word file", type=['docx'], key="file_uploader")
#     if uploaded_file is not None:
#         with st.spinner('Processing the document and generating biography...'):
#             document_text = read_text_from_docx(uploaded_file)
#             biography_text = biography(document_text)
#             st.session_state.biography_text = biography_text  # Store biography text in session state
#             st.session_state.biography_generated = True
#             st.rerun()  # Use st.rerun() instead of experimental_rerun()
# else:
#     st.success("Biography generated successfully.")
    
#     # Add a heading for the biography preview
#     st.markdown("""
#         <h2 style='text-align: left; margin-top: 20px;'>Biography</h2>
#         """, unsafe_allow_html=True)
    
#     st.write(st.session_state.biography_text)
#     doc_io = create_word_document(st.session_state.biography_text)
#     st.download_button(label="Download Biography in Word Format",
#                        data=doc_io,
#                        file_name="biography.docx",
#                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
#     if st.button("Generate new biography"):
#         st.session_state.biography_generated = False
#         st.session_state.biography_text = ""
#         st.rerun()  

# # Add footer to the page
# st.markdown(footer_html, unsafe_allow_html=True)





import streamlit as st
from document_reader_v5 import read_text_from_docx, biography
from docx import Document
import io

# This must be the first Streamlit command used.
st.set_page_config(page_title="German Biography Generator", layout="wide")

# Custom header with HTML and CSS
header_html = """
    <div style="background-color:#618264; padding:10px; border-radius:10px">
    <h1 style="color:white; text-align:center;">German Biography Generator</h1>
    <p style="color:white; text-align:center;">We generate short biographies in German language using interview text of any length. Please upload the interview text in Word format to generate biography of the interviewee.</p>
    </div>
    """
st.markdown(header_html, unsafe_allow_html=True)

# Inject CSS for button color and footer positioning
st.markdown("""
<style>
.stButton>button {
    border: 1px solid #4CAF50;
    border-radius: 4px; /* Match Streamlit's default button radius */
    color: black; /* Match Streamlit's default button text color */
    background-color: #D0E7D2; /* Custom button color */
    font-size: 16px; /* Adjust to match Streamlit's default button size, if needed */
}

footer {
    position: bottom;
    bottom: 0;
    background-color: #333;
    color: #FFFFFF;
    text-align: center;
    padding: 10px 0;
    margin-top: 50px;
    border-radius: 10px;
}

footer a {
    color: #FFFFFF;
    text-decoration: none;
}

footer a:hover {
    text-decoration: underline;
}
</style>
""", unsafe_allow_html=True)

# Footer HTML
footer_html = """
<footer>
    <p style="font-size: 14px;">Contact developer: <a href="mailto:srujanprakashgowda3@gmail.com">srujanprakashgowda3@gmail.com</a></p>
    <p style="font-size: 14px;">
        <a href="https://www.linkedin.com/in/srujanpg/" target="_blank">LinkedIn</a> |
        <a href="https://github.com/SrujanGowda1998" target="_blank">GitHub</a>
    </p>
</footer>
"""

def create_word_document(biography_text):
    doc = Document()
    doc.add_paragraph(biography_text)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def load_sample_document():
    """Load and return the sample document."""
    with open("Transkript_sample_2.docx", "rb") as file:
        return file.read()

# Section for sample document download
st.sidebar.markdown("### Test the App")
st.sidebar.markdown("Don't have an interview transcript??!... We got you... Download a sample interview transcript below and test the app.")
sample_doc_data = load_sample_document()
st.sidebar.download_button(
    label="Download Sample Document",
    data=sample_doc_data,
    file_name="Transkript_sample_2.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

if 'biography_generated' not in st.session_state:
    st.session_state.biography_generated = False

if not st.session_state.biography_generated:
    uploaded_file = st.file_uploader("Choose a Word file", type=['docx'], key="file_uploader")
    if uploaded_file is not None:
        with st.spinner('Processing the document and generating biography...'):
            document_text = read_text_from_docx(uploaded_file)
            biography_text = biography(document_text)
            st.session_state.biography_text = biography_text  # Store biography text in session state
            st.session_state.biography_generated = True
            st.rerun()
else:
    st.success("Biography generated successfully.")
    
    # Add a heading for the biography preview
    st.markdown("""
        <h2 style='text-align: left; margin-top: 20px;'>Biography</h2>
        """, unsafe_allow_html=True)
    
    st.write(st.session_state.biography_text)
    doc_io = create_word_document(st.session_state.biography_text)
    st.download_button(label="Download Biography in Word Format",
                       data=doc_io,
                       file_name="biography.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    if st.button("Generate new biography"):
        st.session_state.biography_generated = False
        st.session_state.biography_text = ""
        st.rerun()

# Add footer to the page
st.markdown(footer_html, unsafe_allow_html=True)
